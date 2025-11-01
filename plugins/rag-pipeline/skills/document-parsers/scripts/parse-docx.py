#!/usr/bin/env python3
"""
Functional DOCX parser with structure preservation
Extracts text, tables, and metadata from Word documents
"""

import argparse
import json
import sys
from pathlib import Path
from typing import Dict, List, Any


class DOCXParser:
    """Word document parser"""

    def __init__(self):
        try:
            from docx import Document
            self.Document = Document
        except ImportError:
            print("Error: python-docx not installed. Run: pip install python-docx")
            sys.exit(1)

    def parse(
        self,
        filepath: str,
        preserve_structure: bool = False,
        extract_tables: bool = True,
        extract_metadata: bool = True
    ) -> Dict[str, Any]:
        """
        Parse DOCX file

        Args:
            filepath: Path to DOCX file
            preserve_structure: Preserve paragraph structure and styles
            extract_tables: Extract tables from document
            extract_metadata: Extract document metadata

        Returns:
            Dictionary with text, tables, and metadata
        """
        doc = self.Document(filepath)

        result = {
            "filepath": filepath,
            "text": "",
            "paragraphs": [],
            "tables": [],
            "metadata": {}
        }

        # Extract metadata
        if extract_metadata:
            core_props = doc.core_properties
            result["metadata"] = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "last_modified_by": core_props.last_modified_by or "",
                "category": core_props.category or "",
                "comments": core_props.comments or ""
            }

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            if preserve_structure:
                para_data = {
                    "text": text,
                    "style": para.style.name if para.style else "Normal"
                }
                result["paragraphs"].append(para_data)

            result["text"] += text + "\n"

        # Extract tables
        if extract_tables:
            for table_idx, table in enumerate(doc.tables):
                table_data = {
                    "table_index": table_idx + 1,
                    "rows": []
                }

                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data["rows"].append(row_data)

                result["tables"].append(table_data)

        return result


def format_text_output(result: Dict[str, Any], tables_only: bool = False) -> str:
    """Format result as plain text"""
    output = ""

    if not tables_only:
        # Add metadata
        if result.get("metadata"):
            metadata = result["metadata"]
            if any(metadata.values()):
                output += "=== Document Metadata ===\n"
                for key, value in metadata.items():
                    if value:
                        output += f"{key.title()}: {value}\n"
                output += "\n"

        # Add content
        if result.get("text"):
            output += "=== Document Content ===\n"
            output += result["text"] + "\n"

    # Add tables
    if result.get("tables"):
        output += "\n=== Tables ===\n\n"
        for table in result["tables"]:
            output += f"Table {table['table_index']}:\n"
            for row in table["rows"]:
                output += " | ".join(row) + "\n"
            output += "\n"

    return output


def format_json_output(result: Dict[str, Any]) -> str:
    """Format result as JSON"""
    return json.dumps(result, indent=2, ensure_ascii=False)


def format_markdown_output(result: Dict[str, Any]) -> str:
    """Format result as Markdown"""
    output = ""

    # Add metadata
    if result.get("metadata"):
        metadata = result["metadata"]
        if any(metadata.values()):
            output += "# Document Metadata\n\n"
            for key, value in metadata.items():
                if value:
                    output += f"**{key.title()}:** {value}  \n"
            output += "\n"

    # Add content with structure
    if result.get("paragraphs"):
        output += "# Content\n\n"
        for para in result["paragraphs"]:
            style = para.get("style", "Normal")
            text = para["text"]

            # Convert heading styles to markdown
            if "Heading 1" in style:
                output += f"# {text}\n\n"
            elif "Heading 2" in style:
                output += f"## {text}\n\n"
            elif "Heading 3" in style:
                output += f"### {text}\n\n"
            elif "Heading 4" in style:
                output += f"#### {text}\n\n"
            else:
                output += f"{text}\n\n"
    elif result.get("text"):
        output += result["text"] + "\n\n"

    # Add tables
    if result.get("tables"):
        output += "# Tables\n\n"
        for table in result["tables"]:
            output += f"## Table {table['table_index']}\n\n"

            # Create markdown table
            rows = table["rows"]
            if rows:
                # Header row
                output += "| " + " | ".join(rows[0]) + " |\n"
                output += "| " + " | ".join(["---"] * len(rows[0])) + " |\n"

                # Data rows
                for row in rows[1:]:
                    output += "| " + " | ".join(row) + " |\n"

                output += "\n"

    return output


def main():
    parser = argparse.ArgumentParser(
        description="Parse DOCX files with structure preservation"
    )
    parser.add_argument(
        "filepath",
        help="Path to DOCX file"
    )
    parser.add_argument(
        "--preserve-structure",
        action="store_true",
        help="Preserve paragraph structure and styles"
    )
    parser.add_argument(
        "--no-tables",
        action="store_true",
        help="Skip table extraction"
    )
    parser.add_argument(
        "--tables-only",
        action="store_true",
        help="Extract only tables, skip text"
    )
    parser.add_argument(
        "--no-metadata",
        action="store_true",
        help="Skip metadata extraction"
    )
    parser.add_argument(
        "--output",
        help="Output file path (default: stdout)"
    )
    parser.add_argument(
        "--format",
        choices=["text", "json", "markdown"],
        default="text",
        help="Output format (default: text)"
    )

    args = parser.parse_args()

    # Check file exists
    if not Path(args.filepath).exists():
        print(f"Error: File not found: {args.filepath}", file=sys.stderr)
        sys.exit(1)

    # Parse DOCX
    try:
        docx_parser = DOCXParser()
        result = docx_parser.parse(
            args.filepath,
            preserve_structure=args.preserve_structure or args.format == "markdown",
            extract_tables=not args.no_tables,
            extract_metadata=not args.no_metadata
        )

        # Format output
        if args.format == "json":
            output = format_json_output(result)
        elif args.format == "markdown":
            output = format_markdown_output(result)
        else:
            output = format_text_output(result, tables_only=args.tables_only)

        # Write output
        if args.output:
            with open(args.output, "w", encoding="utf-8") as f:
                f.write(output)
            print(f"Output written to: {args.output}")
        else:
            print(output)

    except Exception as e:
        print(f"Error parsing DOCX: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
