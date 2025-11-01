#!/usr/bin/env python3
"""
Multi-Format Document Parser Template
Universal parser for PDF, DOCX, HTML, Markdown, and TXT files
"""

import os
import mimetypes
from pathlib import Path
from typing import Dict, List, Optional, Any, Union
from dataclasses import dataclass, field


@dataclass
class ParseResult:
    """Result from parsing a document"""
    filepath: str
    format: str
    text: str
    chunks: List['TextChunk'] = field(default_factory=list)
    tables: List[Dict] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    error: Optional[str] = None


@dataclass
class TextChunk:
    """Chunk of text with metadata"""
    id: str
    text: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    start_index: int = 0
    end_index: int = 0


class MultiFormatParser:
    """
    Universal document parser with automatic format detection

    Supports:
    - PDF (PyPDF2, PDFPlumber, LlamaParse)
    - DOCX (python-docx)
    - HTML (BeautifulSoup)
    - Markdown (markdown)
    - TXT (plain text)

    Features:
    - Automatic chunking for RAG
    - Table extraction
    - Metadata extraction
    - Error handling with fallbacks
    """

    def __init__(
        self,
        llamaparse_api_key: Optional[str] = None,
        use_ocr: bool = False,
        chunk_size: int = 1000,
        chunk_overlap: int = 100,
        prefer_llamaparse: bool = False
    ):
        """
        Initialize parser

        Args:
            llamaparse_api_key: API key for LlamaParse (optional)
            use_ocr: Enable OCR for scanned documents
            chunk_size: Size of text chunks for RAG
            chunk_overlap: Overlap between chunks
            prefer_llamaparse: Use LlamaParse for PDFs when available
        """
        self.llamaparse_api_key = llamaparse_api_key or os.getenv("LLAMA_CLOUD_API_KEY")
        self.use_ocr = use_ocr
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.prefer_llamaparse = prefer_llamaparse and self.llamaparse_api_key

        # Initialize parsers lazily
        self._pdf_parser = None
        self._docx_parser = None
        self._html_parser = None

    def parse_file(self, filepath: Union[str, Path]) -> ParseResult:
        """
        Parse a file with automatic format detection

        Args:
            filepath: Path to file

        Returns:
            ParseResult with extracted content
        """
        filepath = str(filepath)

        if not os.path.exists(filepath):
            return ParseResult(
                filepath=filepath,
                format="unknown",
                text="",
                error=f"File not found: {filepath}"
            )

        # Detect format
        file_format = self._detect_format(filepath)

        try:
            # Parse based on format
            if file_format == "pdf":
                result = self._parse_pdf(filepath)
            elif file_format == "docx":
                result = self._parse_docx(filepath)
            elif file_format == "html":
                result = self._parse_html(filepath)
            elif file_format == "markdown":
                result = self._parse_markdown(filepath)
            elif file_format == "txt":
                result = self._parse_txt(filepath)
            else:
                return ParseResult(
                    filepath=filepath,
                    format=file_format,
                    text="",
                    error=f"Unsupported format: {file_format}"
                )

            # Add chunks
            if result.text:
                result.chunks = self._create_chunks(result.text, result.metadata)

            return result

        except Exception as e:
            return ParseResult(
                filepath=filepath,
                format=file_format,
                text="",
                error=str(e)
            )

    def parse_directory(
        self,
        directory: Union[str, Path],
        recursive: bool = True,
        extensions: Optional[List[str]] = None
    ) -> Dict[str, ParseResult]:
        """
        Parse all documents in a directory

        Args:
            directory: Path to directory
            recursive: Search subdirectories
            extensions: Filter by extensions (e.g., ['.pdf', '.docx'])

        Returns:
            Dictionary mapping filepath to ParseResult
        """
        directory = Path(directory)
        results = {}

        # Default extensions
        if extensions is None:
            extensions = ['.pdf', '.docx', '.html', '.htm', '.md', '.txt']

        # Find files
        pattern = "**/*" if recursive else "*"
        for filepath in directory.glob(pattern):
            if filepath.is_file() and filepath.suffix.lower() in extensions:
                results[str(filepath)] = self.parse_file(filepath)

        return results

    def _detect_format(self, filepath: str) -> str:
        """Detect file format from extension and MIME type"""
        ext = Path(filepath).suffix.lower()

        # Check extension first
        ext_map = {
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.doc': 'doc',  # Old Word format (not supported)
            '.html': 'html',
            '.htm': 'html',
            '.md': 'markdown',
            '.markdown': 'markdown',
            '.txt': 'txt'
        }

        if ext in ext_map:
            return ext_map[ext]

        # Try MIME type
        mime_type, _ = mimetypes.guess_type(filepath)
        if mime_type:
            if 'pdf' in mime_type:
                return 'pdf'
            elif 'word' in mime_type:
                return 'docx'
            elif 'html' in mime_type:
                return 'html'

        return 'txt'  # Default to text

    def _parse_pdf(self, filepath: str) -> ParseResult:
        """Parse PDF file"""
        if self.prefer_llamaparse and self.llamaparse_api_key:
            # Try LlamaParse first
            try:
                from llama_parse import LlamaParse

                parser = LlamaParse(
                    api_key=self.llamaparse_api_key,
                    result_type="markdown",
                    verbose=False
                )

                documents = parser.load_data(filepath)
                text = "\n\n".join(doc.text for doc in documents)

                return ParseResult(
                    filepath=filepath,
                    format="pdf",
                    text=text,
                    metadata={"parser": "llamaparse"}
                )
            except Exception as e:
                print(f"LlamaParse failed, falling back to local parser: {e}")

        # Fallback to local parsers
        try:
            import pdfplumber

            result = ParseResult(filepath=filepath, format="pdf", metadata={"parser": "pdfplumber"})

            with pdfplumber.open(filepath) as pdf:
                # Extract metadata
                if pdf.metadata:
                    result.metadata.update({
                        k.replace("/", ""): v for k, v in pdf.metadata.items()
                    })

                # Extract text and tables
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    text_parts.append(page_text)

                    # Extract tables
                    tables = page.extract_tables()
                    if tables:
                        for idx, table in enumerate(tables):
                            result.tables.append({
                                "page": page.page_number,
                                "table_index": idx,
                                "data": table
                            })

                result.text = "\n\n".join(text_parts)

            return result

        except ImportError:
            # Final fallback to PyPDF2
            from PyPDF2 import PdfReader

            reader = PdfReader(filepath)
            text_parts = []

            for page in reader.pages:
                text_parts.append(page.extract_text())

            return ParseResult(
                filepath=filepath,
                format="pdf",
                text="\n\n".join(text_parts),
                metadata={"parser": "pypdf2"}
            )

    def _parse_docx(self, filepath: str) -> ParseResult:
        """Parse DOCX file"""
        from docx import Document

        doc = Document(filepath)
        result = ParseResult(filepath=filepath, format="docx")

        # Extract metadata
        core_props = doc.core_properties
        result.metadata = {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "created": str(core_props.created) if core_props.created else "",
        }

        # Extract text
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        result.text = "\n".join(text_parts)

        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                table_data.append([cell.text.strip() for cell in row.cells])
            result.tables.append({
                "table_index": table_idx + 1,
                "data": table_data
            })

        return result

    def _parse_html(self, filepath: str) -> ParseResult:
        """Parse HTML file"""
        from bs4 import BeautifulSoup

        with open(filepath, 'r', encoding='utf-8') as f:
            html_content = f.read()

        soup = BeautifulSoup(html_content, 'lxml')

        result = ParseResult(filepath=filepath, format="html")

        # Extract metadata
        if soup.title:
            result.metadata["title"] = soup.title.string.strip()

        for meta in soup.find_all("meta"):
            name = meta.get("name") or meta.get("property")
            content = meta.get("content")
            if name and content:
                result.metadata[name] = content

        # Remove unwanted elements
        for element in soup(["script", "style", "noscript"]):
            element.decompose()

        # Extract clean text
        result.text = soup.get_text(separator="\n", strip=True)

        # Clean up whitespace
        lines = [line.strip() for line in result.text.split("\n") if line.strip()]
        result.text = "\n".join(lines)

        return result

    def _parse_markdown(self, filepath: str) -> ParseResult:
        """Parse Markdown file"""
        with open(filepath, 'r', encoding='utf-8') as f:
            text = f.read()

        # Extract title from first heading
        metadata = {}
        lines = text.split("\n")
        if lines and lines[0].startswith("#"):
            metadata["title"] = lines[0].lstrip("#").strip()

        return ParseResult(
            filepath=filepath,
            format="markdown",
            text=text,
            metadata=metadata
        )

    def _parse_txt(self, filepath: str) -> ParseResult:
        """Parse plain text file"""
        with open(filepath, 'r', encoding='utf-8') as f:
            text = f.read()

        return ParseResult(
            filepath=filepath,
            format="txt",
            text=text
        )

    def _create_chunks(self, text: str, metadata: Dict[str, Any]) -> List[TextChunk]:
        """Create text chunks for RAG"""
        chunks = []
        words = text.split()

        # Simple word-based chunking
        chunk_id = 0
        start_idx = 0

        while start_idx < len(words):
            end_idx = min(start_idx + self.chunk_size, len(words))
            chunk_words = words[start_idx:end_idx]
            chunk_text = " ".join(chunk_words)

            chunks.append(TextChunk(
                id=f"chunk_{chunk_id}",
                text=chunk_text,
                metadata=metadata.copy(),
                start_index=start_idx,
                end_index=end_idx
            ))

            chunk_id += 1
            start_idx = end_idx - self.chunk_overlap

        return chunks


# Example usage
if __name__ == "__main__":
    # Initialize parser
    parser = MultiFormatParser(
        chunk_size=512,
        chunk_overlap=50,
        prefer_llamaparse=False
    )

    # Parse single file
    result = parser.parse_file("document.pdf")
    if result.error:
        print(f"Error: {result.error}")
    else:
        print(f"Parsed {result.format} file: {result.filepath}")
        print(f"Text length: {len(result.text)} characters")
        print(f"Chunks: {len(result.chunks)}")
        print(f"Tables: {len(result.tables)}")
        print(f"Metadata: {result.metadata}")

    # Parse directory
    results = parser.parse_directory("./documents/", recursive=True)
    for filepath, result in results.items():
        if result.error:
            print(f"✗ {filepath}: {result.error}")
        else:
            print(f"✓ {filepath}: {len(result.chunks)} chunks")
